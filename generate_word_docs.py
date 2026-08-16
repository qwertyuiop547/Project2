import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def create_document():
    doc = Document()
    
    # 1-inch standard margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Color Palette
    NAVY = RGBColor(15, 23, 42)       # #0f172a
    EMERALD = RGBColor(5, 150, 105)   # #059669
    DARK_BLUE = RGBColor(30, 58, 138) # #1e3a8a
    GRAY = RGBColor(100, 116, 139)    # #64748b
    TEXT_DARK = RGBColor(30, 41, 59)  # #1e293b

    # Base typography
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Segoe UI'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = TEXT_DARK

    # --- TITLE / COVER HEADER ---
    p_kicker = doc.add_paragraph()
    r_kicker = p_kicker.add_run("REPUBLIC OF THE PHILIPPINES • PROVINCE OF RIZAL • BARANGAY BURGOS")
    r_kicker.font.size = Pt(9)
    r_kicker.font.bold = True
    r_kicker.font.color.rgb = EMERALD
    p_kicker.paragraph_format.space_after = Pt(2)

    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Barangay Burgos e-Services & AI Governance Portal")
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = NAVY
    p_title.paragraph_format.space_after = Pt(2)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("End-to-End Technical & System Documentation Manual (Full System Specification)")
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = DARK_BLUE
    p_sub.paragraph_format.space_after = Pt(12)

    # Metadata Summary Box
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    meta_items = [
        ("System Name:", "Barangay Burgos Digital Governance & AI Portal"),
        ("Live Cloud URL:", "https://barangay-portal-jqxc.onrender.com/"),
        ("Source Code Repository:", "https://github.com/qwertyuiop547/Project2.git"),
        ("Documentation Version:", "Version 3.0 (Comprehensive Full System Release)")
    ]

    for idx, (label, val) in enumerate(meta_items):
        row = meta_table.rows[idx]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = val
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        for cell in row.cells:
            set_cell_background(cell, "F1F5F9")
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Section Heading Utilities
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = NAVY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(11)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = EMERALD
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.bold = True
        r.font.color.rgb = DARK_BLUE
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        p.paragraph_format.space_after = Pt(2.5)
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.color.rgb = NAVY
        r2 = p.add_run(text)
        r2.font.color.rgb = TEXT_DARK

    def add_callout(title, text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_background(cell, "ECFDF5")
        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
        p = cell.paragraphs[0]
        r1 = p.add_run(title + "\n")
        r1.font.bold = True
        r1.font.color.rgb = EMERALD
        r1.font.size = Pt(9.5)
        r2 = p.add_run(text)
        r2.font.size = Pt(9)

    # =========================================================================
    # SECTION 1: SYSTEM OVERVIEW & ARCHITECTURE
    # =========================================================================
    add_h1("1. System Overview & Problem Statement")
    doc.add_paragraph(
        "In traditional local government operations across the Philippines, barangay transactions require physical in-person visits, "
        "manual paper logbooks, and protracted waiting times. Complaints are often filed verbally or on loose paper, making tracking, "
        "status updates, and dispute resolution difficult to measure and audit."
    )
    doc.add_paragraph(
        "The Barangay Burgos e-Services & AI Governance Portal solves this by providing a unified, full-stack cloud platform. "
        "It empowers citizens to file verified complaints, request documents, voice community ideas, and converse with an AI Assistant 24/7. "
        "Simultaneously, it equips barangay officials with analytics, automated case workflows, and certificate issuance tools."
    )

    add_h2("1.1 Full-Stack Technical Stack Architecture")
    add_bullet("Frontend SPA Layer: ", "React 18.2 with Vite 5 build toolchain, Zustand for client state management, TanStack React Query for cached server state synchronization, Lucide React iconography, and Vanilla CSS design tokens.")
    add_bullet("Backend API Layer: ", "Node.js (v18+) with Express 4 REST architecture, Express Validator for payload sanitization, JSON Web Token (JWT) stateless auth, and Server-Sent Events (SSE) for real-time streaming.")
    add_bullet("Data Persistence Layer: ", "PostgreSQL relational database hosted in cloud infrastructure, managed through Prisma 5 ORM with declarative schema migrations and seed scripts.")
    add_bullet("AI Intelligence Layer: ", "Hybrid AI router supporting Local Open-Source LLMs (Ollama Llama 3.2), Cloud APIs (Google Gemini, Groq, OpenAI), and a zero-dependency Tagalog Rule-Based NLP engine.")

    # =========================================================================
    # SECTION 2: PROJECT DIRECTORY STRUCTURE
    # =========================================================================
    add_h1("2. Complete Project Directory Structure")
    doc.add_paragraph("The codebase is organized as a full-stack monorepo with clear separation between client, server, and database:")

    dir_table = doc.add_table(rows=9, cols=2)
    dir_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    dir_headers = ["Path / Directory", "Component Responsibility"]
    for i, title in enumerate(dir_headers):
        cell = dir_table.rows[0].cells[i]
        cell.text = title
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(9)

    dir_data = [
        ("client/src/pages/", "Application pages (Home, Login, Register, Dashboard, Complaints, Suggestions, Services, Announcements)."),
        ("client/src/components/", "Reusable UI components (Navbar, Layout, BarangayChatbot, AccessibilityPanel, FloatingCTA)."),
        ("client/src/context/", "Global state providers (AuthContext for user sessions and permissions)."),
        ("client/src/lib/", "Client utilities, Axios API client interceptors, and chatStream SSE streaming reader."),
        ("server/src/routes/", "Express route controllers (auth.js, ai.js, complaints.js, services.js, suggestions.js, announcements.js)."),
        ("server/src/services/ai.js", "Core AI engine: Ollama integration, Gemini API client, NLP regex dictionaries, and rule fallbacks."),
        ("server/src/middleware/", "Authentication middleware (auth.js for JWT verification and role checks)."),
        ("server/prisma/", "Database schema definition (schema.prisma) and pre-seeded accounts and categories (seed.js).")
    ]

    for row_idx, (path_val, desc_val) in enumerate(dir_data, start=1):
        row = dir_table.rows[row_idx]
        row.cells[0].text = path_val
        row.cells[1].text = desc_val
        for col_idx, cell in enumerate(row.cells):
            set_cell_background(cell, "FFFFFF" if row_idx % 2 == 1 else "F8FAFC")
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(8.5)
            if col_idx == 0:
                p.runs[0].font.bold = True

    # =========================================================================
    # SECTION 3: ROLE-BASED ACCESS CONTROL (RBAC)
    # =========================================================================
    add_h1("3. Role-Based Access Control & User Personas")
    doc.add_paragraph("The system implements strict permission isolation across three primary roles:")

    role_table = doc.add_table(rows=4, cols=3)
    role_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    role_headers = ["User Role", "Target Persona", "System Capabilities & Permissions"]
    for i, title in enumerate(role_headers):
        cell = role_table.rows[0].cells[i]
        cell.text = title
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(9)

    role_data = [
        ("RESIDENT", "Citizens / Constituents", "File complaints with AI assistance, request barangay documents, submit and upvote community suggestions, chat with AI Assistant, track live filing progress."),
        ("SECRETARY", "Barangay Secretary & Admin", "Review and process incoming complaints, update case statuses, issue official certificates, manage indigent records, publish official announcements."),
        ("CHAIRMAN", "Punong Barangay (Captain)", "Full municipal governance analytics, review case resolution metrics, oversee Lupon Tagapamayapa dispute mediation, and manage barangay operations.")
    ]

    for row_idx, (r_name, r_target, r_perms) in enumerate(role_data, start=1):
        row = role_table.rows[row_idx]
        row.cells[0].text = r_name
        row.cells[1].text = r_target
        row.cells[2].text = r_perms
        for col_idx, cell in enumerate(row.cells):
            set_cell_background(cell, "FFFFFF" if row_idx % 2 == 1 else "F8FAFC")
            set_cell_margins(cell, top=70, bottom=70, left=100, right=100)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(8.5)
            if col_idx == 0:
                p.runs[0].font.bold = True

    add_bullet("Pre-Seeded Resident: ", "resident@example.com | password123 (Name: Pedro Reyes)")
    add_bullet("Pre-Seeded Secretary: ", "secretary@barangay.gov.ph | password123 (Name: Maria Santos)")
    add_bullet("Pre-Seeded Chairman: ", "chairman@barangay.gov.ph | password123 (Name: Juan Dela Cruz)")

    # =========================================================================
    # SECTION 4: AI FEATURES DEEP DIVE
    # =========================================================================
    add_h1("4. 🤖 Comprehensive AI Intelligence Subsystem")
    doc.add_paragraph(
        "The Artificial Intelligence layer is one of the central pillars of the portal. It is specifically calibrated for Philippine "
        "local government contexts, handling multilingual input (Tagalog, English, Taglish) and guaranteeing zero failure."
    )

    add_h2("4.1 Multi-Provider AI Routing Pipeline")
    doc.add_paragraph(
        "Located in 'server/src/services/ai.js', the AI engine evaluates provider availability dynamically:"
    )
    add_bullet("1. Local Ollama LLM (Preferred on-premise): ", "Queries a local Ollama server hosting models such as Llama 3.2 (1B/3B) or Mistral. Zero external API fees, high performance, and 100% data confidentiality.")
    add_bullet("2. Google Gemini API (Cloud): ", "Cloud LLM integration activated when GEMINI_API_KEY is defined in environment variables.")
    add_bullet("3. Groq / OpenAI API: ", "Ultra-fast alternative cloud providers for high-concurrency environments.")
    add_bullet("4. Smart Assist NLP Engine (Deterministic Fallback): ", "When no LLM or internet connectivity is present, the built-in rule engine autonomously parses keywords, scores hazards, and synthesizes structured Tagalog complaints.")

    add_h2("4.2 AI Complaint Assistant & Auto-Refinement")
    doc.add_paragraph(
        "When a citizen submits a complaint, the AI performs three automated operations:"
    )
    add_bullet("1. Category Classification: ", "Matches keywords to determine if the report belongs to Infrastructure, Sanitation, Public Safety, Noise & Disturbance, or Others.")
    add_bullet("2. Hazard & Priority Scoring: ", "Analyzes urgency markers (e.g. 'sunog', 'baha', 'aksidente', 'nakawan', 'poste') to assign URGENT, HIGH, MEDIUM, or LOW priority.")
    add_bullet("3. Formal Tagalog Synthesis: ", "Polishes colloquial or emotional citizen input into courteous, respectful Tagalog suitable for barangay records while strictly preserving landmarks, dates, and locations.")

    add_callout(
        "💡 AI Complaint Transformation Example:",
        "• Raw Citizen Input: \"may butas kalsada tapat ng tindahan ni aling nena delikado sa motor baha pa\"\n"
        "• Polished Report Title: \"Lubak at Pagbaha sa Kalsada sa Purok 3\"\n"
        "• Polished Description: \"Nais kong i-report ang malaking butas sa kalsada sa tapat ng Nena Store na nagdudulot ng panganib sa mga nagmomotor at nagiging sanhi ng pagbaha. Humihiling po kami ng agarang aksyon mula sa barangay.\"\n"
        "• Classification: Infrastructure | Priority: HIGH"
    )

    add_h2("4.3 Barangay AI Chatbot (Real-Time SSE Streaming)")
    doc.add_paragraph(
        "The Barangay Chatbot provides an interactive virtual desk accessible from any page:"
    )
    add_bullet("Real-Time Token Streaming: ", "Uses Server-Sent Events (SSE) via '/api/ai/chat/stream' for instantaneous typewriter-style responses.")
    add_bullet("Dynamic Portal Context: ", "Injects real-time barangay services, processing fees, document requirements, and the resident's active complaints into the system prompt.")
    add_bullet("Instant FAQ Acceleration: ", "Frequently asked questions (e.g. office hours, clearances, hotlines) are answered via fast-path regex in < 20 milliseconds.")

    # =========================================================================
    # SECTION 5: CORE CIVIC MODULES
    # =========================================================================
    add_h1("5. Core Civic Modules & Features")
    
    add_h2("5.1 Civic Homepage ('/')")
    add_bullet("Live PST Clock: ", "Real-time Philippine Standard Time clock displayed at the top bar.")
    add_bullet("Time-Aware Greeting: ", "Dynamic greetings ('Magandang Umaga/Hapon/Gabi') with live service status pulse indicator.")
    add_bullet("Civic Stats Ribbon: ", "Highlights 4 transparency metrics (24/7 Digital Filing, 24-48h Processing, 98.5% Resolution Rate, 3,500+ Residents).")
    add_bullet("Services Explorer: ", "Searchable directory with interactive filter tabs (All, Clearance & Docs, Assistance, Business, Lupon Tagapamayapa).")
    add_bullet("Emergency Hotlines Bar: ", "One-click dialer for Barangay Desk, Tanod Patrol, Health Center, and Fire/PNP.")
    add_bullet("Interactive FAQ Accordion: ", "Expandable answers for common community inquiries.")

    add_h2("5.2 Authentication & Security ('/login', '/register')")
    add_bullet("Quick Demo Login: ", "1-click autofill buttons for Resident, Chairman, and Secretary testing.")
    add_bullet("Animated Password Strength Meter: ", "Color-coded visual indicator assessing password complexity in real time.")
    add_bullet("Data Privacy Act Compliance (R.A. 10173): ", "Mandatory consent checkbox protecting citizen personal data.")

    add_h2("5.3 Resident Dashboard ('/dashboard')")
    add_bullet("Live Metric Cards: ", "Summary badges for Pending, In Progress, Resolved complaints, and submitted ideas.")
    add_bullet("Quick Action Cards: ", "Direct links to File Complaint, Request Document, Submit Idea, and Read Announcements.")
    add_bullet("My Requests Table: ", "Comprehensive tracking table showing filing date, category, status, and tracking ID.")

    add_h2("5.4 Official Governance Dashboard ('/dashboard')")
    add_bullet("5-Column KPI Analytics Ribbon: ", "Tracks Total Complaints, For Action Queue, Resolved Rate, Pending Clearances, and Ideas.")
    add_bullet("Dispute & Case Workflow: ", "Allows officials to transition cases from PENDING ➔ IN_PROGRESS ➔ RESOLVED / DISMISSED with official notes.")
    add_bullet("Interactive Filters & Search: ", "Filter table by category, status, or search keywords.")

    add_h2("5.5 Document & Certificate Processing ('/services')")
    add_bullet("Available Services: ", "Barangay Clearance, Certificate of Indigency, Barangay Business Clearance, Solo Parent Certificate, and Barangay ID.")
    add_bullet("Transparent Fee & Requirements: ", "Displays exact processing fees, estimated release days, and required valid IDs.")

    add_h2("5.6 Community Voice & Suggestions ('/suggestions')")
    add_bullet("Idea Crowdsourcing: ", "Residents submit community improvement proposals with category tags.")
    add_bullet("Community Upvoting: ", "Constituents upvote ideas, highlighting popular initiatives to barangay leadership.")

    # =========================================================================
    # SECTION 6: ACCESSIBILITY (WCAG AAA)
    # =========================================================================
    add_h1("6. Accessibility & Universal Inclusive Design")
    doc.add_paragraph(
        "To ensure usability for all citizens, including senior citizens and visually impaired users, the portal complies with WCAG 2.1 AAA:"
    )
    add_bullet("High Contrast OLED Mode: ", "Replaces light colors with pure black (#000000) background, bold white (#ffffff) typography, high-visibility yellow (#ffff00) action buttons, and neon green (#00ff88) accents.")
    add_bullet("Dynamic Font Scaling: ", "Users can switch between Normal (100%), Large (112%), and Extra Large (124%) font sizes across all screens.")
    add_bullet("Full Keyboard Navigation: ", "Semantic HTML landmarks, visible focus rings, and screen-reader accessible ARIA attributes.")

    # =========================================================================
    # SECTION 7: DATABASE SCHEMA & DATA DICTIONARY
    # =========================================================================
    add_h1("7. Database Schema & Data Dictionary (Prisma ORM)")
    doc.add_paragraph("The database schema is defined in 'server/prisma/schema.prisma' with the following core entities:")

    db_table = doc.add_table(rows=7, cols=3)
    db_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    db_headers = ["Entity / Model", "Key Attributes & Fields", "Relationships & Purpose"]
    for i, title in enumerate(db_headers):
        cell = db_table.rows[0].cells[i]
        cell.text = title
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(9)

    db_data = [
        ("User", "id, email, password, firstName, lastName, phone, address, role (RESIDENT / SECRETARY / CHAIRMAN), isApproved", "Stores citizen and official user accounts with encrypted passwords."),
        ("Complaint", "id, title, description, location, priority (LOW/MED/HIGH/URGENT), status (PENDING/IN_PROGRESS/RESOLVED/DISMISSED), userId, categoryId", "Stores all citizen incident reports and official status timeline."),
        ("ComplaintCategory", "id, name, description, icon", "Defines complaint categories (Infrastructure, Sanitation, Public Safety, Noise, Others)."),
        ("Service", "id, name, description, fee, processingDays, requirements", "Catalog of barangay certificates, clearances, and document services."),
        ("Suggestion", "id, title, description, category, status, upvotesCount, userId", "Crowdsourced citizen proposals for barangay improvements."),
        ("Announcement", "id, title, content, isUrgent, publishedAt, authorId", "Official barangay bulletins, public advisories, and emergency alerts.")
    ]

    for row_idx, (m_name, m_attr, m_rel) in enumerate(db_data, start=1):
        row = db_table.rows[row_idx]
        row.cells[0].text = m_name
        row.cells[1].text = m_attr
        row.cells[2].text = m_rel
        for col_idx, cell in enumerate(row.cells):
            set_cell_background(cell, "FFFFFF" if row_idx % 2 == 1 else "F8FAFC")
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(8.5)
            if col_idx == 0:
                p.runs[0].font.bold = True

    # =========================================================================
    # SECTION 8: COMPLETE REST API SPECIFICATION
    # =========================================================================
    add_h1("8. Complete REST API Specification")
    
    api_table = doc.add_table(rows=11, cols=4)
    api_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    api_headers = ["HTTP Method", "Endpoint URI", "Function & Expected Behavior", "Auth Required"]
    for i, title in enumerate(api_headers):
        cell = api_table.rows[0].cells[i]
        cell.text = title
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(9)

    api_data = [
        ("POST", "/api/auth/register", "Register new resident account with hashed password", "No"),
        ("POST", "/api/auth/login", "Authenticate user & issue 7-day JWT session token", "No"),
        ("GET", "/api/auth/me", "Fetch profile and role data of authenticated user", "Yes (JWT)"),
        ("GET", "/api/ai/status", "Check active AI provider (Ollama, Gemini, Groq, Smart)", "Yes"),
        ("POST", "/api/ai/complaint-assist", "AI formal rewrite, classification, and priority scoring", "Yes (Resident)"),
        ("POST", "/api/ai/chat", "Send user query to Barangay AI Assistant (JSON)", "Yes"),
        ("POST", "/api/ai/chat/stream", "Stream real-time AI response tokens via SSE", "Yes"),
        ("GET / POST", "/api/complaints", "Retrieve complaint history or submit new complaint", "Yes"),
        ("PATCH", "/api/complaints/:id/status", "Update complaint status & official notes (Officials only)", "Yes (Admin)"),
        ("GET", "/api/dashboard/stats", "Retrieve KPI analytics and summary metrics for dashboard", "Yes")
    ]

    for row_idx, (m_val, ep_val, desc_val, auth_val) in enumerate(api_data, start=1):
        row = api_table.rows[row_idx]
        row.cells[0].text = m_val
        row.cells[1].text = ep_val
        row.cells[2].text = desc_val
        row.cells[3].text = auth_val
        for col_idx, cell in enumerate(row.cells):
            set_cell_background(cell, "FFFFFF" if row_idx % 2 == 1 else "F8FAFC")
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(8.5)
            if col_idx == 0:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = EMERALD

    # =========================================================================
    # SECTION 9: DEPLOYMENT & ENVIRONMENT CONFIGURATION
    # =========================================================================
    add_h1("9. Setup, Environment Variables & Cloud Deployment")
    
    add_h2("9.1 Environment Variables ('server/.env')")
    add_bullet("NODE_ENV: ", "'production' or 'development'")
    add_bullet("PORT: ", "5000 (Local default) or dynamically assigned by cloud host")
    add_bullet("DATABASE_URL: ", "PostgreSQL connection string (e.g. postgresql://user:password@host:5432/dbname)")
    add_bullet("JWT_SECRET: ", "256-bit cryptographically secure secret key for token signing")
    add_bullet("JWT_EXPIRES_IN: ", "'7d' (Token session duration)")
    add_bullet("CORS_ORIGIN: ", "'*' (Production CORS configuration)")
    add_bullet("GEMINI_API_KEY: ", "(Optional) Google Gemini generative API key")
    add_bullet("OLLAMA_BASE_URL: ", "(Optional) Local Ollama server address (default: http://localhost:11434)")

    add_h2("9.2 Production Build & Deployment Commands")
    add_bullet("Install Monorepo Dependencies: ", "npm install")
    add_bullet("Build Frontend Bundle: ", "npm run build (Outputs optimized SPA bundle to client/dist)")
    add_bullet("Sync Database Schema: ", "cd server && npx prisma db push --accept-data-loss")
    add_bullet("Seed Default Accounts: ", "node prisma/seed.js")
    add_bullet("Start Production Server: ", "npm start")

    # Document Footer
    doc.add_paragraph().paragraph_format.space_before = Pt(24)
    p_foot = doc.add_paragraph()
    r_foot = p_foot.add_run("Barangay Burgos Digital Governance System • Full System Documentation • 2026")
    r_foot.font.size = Pt(9)
    r_foot.font.color.rgb = GRAY
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_path = "c:\\Users\\alice\\OneDrive\\Documents\\Project2-main\\BARANGAY_PORTAL_DOCUMENTATION.docx"
    doc.save(output_path)
    print(f"[SUCCESS] Generated Complete Full-System Word document at: {output_path}")

if __name__ == "__main__":
    create_document()
